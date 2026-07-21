"""Юнит-тесты сервиса редактируемого DOCX-шаблона протокола (часть 2).

Фикстурный .docx собирается прямо в тесте через python-docx: docxtpl считает
такой файл валидным Jinja2-шаблоном (плейсхолдеры вписаны в текст параграфов)."""

from __future__ import annotations

from pathlib import Path
import tempfile

import pytest
from docx import Document

from app.models import Protocol, Task
from app.services import protocol_template


def _make_template_docx(path: Path, text: list[str] | None = None) -> Path:
    lines = text or [
        "Протокол {{ title }} от {{ date }} № {{ number }}",
        "{{ body }}",
        "{% for t in tasks %}{{ t.assignment }}|{{ t.responsible }}|{{ t.department }}"
        "|{{ t.deadline }}|{{ t.status }}{% endfor %}",
    ]
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(str(path))
    return path


def _docx_text(path: Path) -> str:
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def _sample_protocol() -> Protocol:
    return Protocol(
        id="proto-1",
        title="Планёрка",
        date="01.02.2030",
        number="9",
        body="Обсудили бюджет.",
        tasks=[
            Task(assignment="Сделать X", responsible="Иванов", department="ИТ",
                 deadline="завтра", status="Новое"),
        ],
    )


def test_extract_placeholders_finds_canonical_and_ignores_loop_var(tmp_path):
    path = _make_template_docx(tmp_path / "tpl.docx")
    found = protocol_template.extract_placeholders(str(path))
    # Ровно канонические поля; переменная цикла t (declared) не попадает.
    assert set(found) == {"title", "date", "number", "body", "tasks"}


def test_auto_map_maps_only_exact_matches():
    mapping = protocol_template.auto_map(["title", "date", "weird", "tasks"])
    assert mapping == {"title": "title", "date": "date", "tasks": "tasks"}
    # Несовпавшие канонические поля просто отсутствуют.
    assert "number" not in mapping
    assert "weird" not in mapping


def test_build_context_shapes_tasks():
    ctx = protocol_template.build_context(_sample_protocol())
    assert ctx["title"] == "Планёрка"
    assert ctx["date"] == "01.02.2030"
    assert ctx["number"] == "9"
    assert ctx["body"] == "Обсудили бюджет."
    assert ctx["tasks"] == [
        {"assignment": "Сделать X", "responsible": "Иванов", "department": "ИТ",
         "deadline": "завтра", "status": "Новое"},
    ]


def test_render_fills_placeholders_with_real_data(tmp_path):
    import json

    from app.models import ProtocolTemplate

    src = _make_template_docx(tmp_path / "tpl.docx")
    tpl = ProtocolTemplate(
        id="t1", docx_path=str(src),
        field_mapping_json=json.dumps(protocol_template.auto_map(
            protocol_template.extract_placeholders(str(src)))),
    )
    out = protocol_template.render(_sample_protocol(), tpl)
    text = _docx_text(out)
    assert "Планёрка" in text
    assert "01.02.2030" in text
    assert "№ 9" in text
    assert "Обсудили бюджет." in text
    # Строка цикла tasks отрендерилась с данными поручения.
    assert "Сделать X|Иванов|ИТ|завтра|Новое" in text


def test_render_honors_custom_mapping(tmp_path):
    """field_mapping управляет тем, какое каноническое поле идёт в плейсхолдер."""
    import json

    from app.models import ProtocolTemplate

    src = _make_template_docx(tmp_path / "tpl.docx", text=["{{ heading }}"])
    # Плейсхолдер файла — heading; маппим его на каноническое поле title.
    tpl = ProtocolTemplate(
        id="t2", docx_path=str(src),
        field_mapping_json=json.dumps({"title": "heading"}),
    )
    out = protocol_template.render(_sample_protocol(), tpl)
    assert "Планёрка" in _docx_text(out)


def test_extract_placeholders_broken_docx_raises_template_error():
    # Файл с расширением .docx, но это не zip -> понятная ошибка, а не сырое исключение.
    bad = Path(tempfile.mkdtemp()) / "broken.docx"
    bad.write_text("это не docx", encoding="utf-8")
    with pytest.raises(protocol_template.ProtocolTemplateError):
        protocol_template.extract_placeholders(str(bad))
