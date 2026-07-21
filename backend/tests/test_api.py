"""Smoke-тесты API через TestClient (с lifespan для создания таблиц)."""

from __future__ import annotations

from datetime import datetime, timezone

import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.db import SessionLocal
from app.models import Protocol, Task


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as c:
        yield c


# --- Хелперы для DOCX-шаблонов протокола (часть 2) ---

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _template_bytes(lines: list[str]) -> bytes:
    """Собрать .docx-шаблон из строк-параграфов (docxtpl считает такой файл валидным)."""
    import io

    from docx import Document

    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_text_from_bytes(content: bytes) -> str:
    import io

    from docx import Document

    return "\n".join(p.text for p in Document(io.BytesIO(content)).paragraphs)


def test_health(client):
    data = client.get("/api/health").json()
    assert data["status"] == "ok"
    assert "ffmpeg" in data and "asr_model" in data


def test_empty_collections(client):
    assert client.get("/api/transcriptions").json() == []
    assert client.get("/api/tasks").json() == []
    lib = client.get("/api/library").json()
    assert lib == {"protocols": [], "transcriptions": []}


def test_search_local_fallback(client):
    # Без Dify уходит в локальный фоллбэк, не падает.
    res = client.post("/api/search", json={"query": "бюджет"})
    assert res.status_code == 200
    assert "hits" in res.json()


def test_upload_validation_rejects_bad_ext(client):
    res = client.post(
        "/api/transcriptions",
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )
    assert res.status_code == 415


def test_transcription_create_and_delete(client):
    res = client.post(
        "/api/transcriptions",
        files={"file": ("clip.mp3", b"\x00\x01\x02\x03", "audio/mpeg")},
    )
    assert res.status_code == 200
    tid = res.json()["id"]

    # запись появилась в списке
    ids = [t["id"] for t in client.get("/api/transcriptions").json()]
    assert tid in ids

    # удаление
    assert client.delete(f"/api/transcriptions/{tid}").status_code == 200
    assert client.get(f"/api/transcriptions/{tid}").status_code == 404


def test_export_missing_object_404(client):
    res = client.post("/api/export", json={"object_type": "protocol", "object_id": "x", "fmt": "md"})
    assert res.status_code == 404


def test_spa_blocks_path_traversal(client):
    # Процентно-закодированный «../» не должен выдавать файлы вне frontend/dist
    # (например config.yaml с ключами). Ответ — 404 или SPA index.html, но
    # НИКОГДА содержимое файла за пределами dist.
    res = client.get("/%2e%2e/%2e%2e/backend/requirements.txt")
    assert "fastapi" not in res.text.lower()


def test_spa_serves_index_for_unknown_route(client):
    # Легитимный роут React Router -> отдаём index.html (200, HTML).
    res = client.get("/library")
    assert res.status_code == 200
    assert "<!doctype html>" in res.text.lower() or '<div id="root"' in res.text


def test_unknown_api_route_returns_404(client):
    # Неизвестные /api/* не проваливаются в SPA-фоллбэк (HTML со статусом 200).
    res = client.get("/api/does-not-exist")
    assert res.status_code == 404


def test_task_update_rejects_invalid_status(client):
    with SessionLocal() as db:
        protocol = Protocol(title="Статусы")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        task = Task(protocol_id=protocol.id, assignment="Задача")
        db.add(task)
        db.commit()
        db.refresh(task)
        task_id = task.id

    # Произвольная строка статуса отклоняется схемой.
    res = client.patch(f"/api/tasks/{task_id}", json={"status": "Абракадабра"})
    assert res.status_code == 422

    # Валидный статус проходит.
    res = client.patch(f"/api/tasks/{task_id}", json={"status": "Выполнено"})
    assert res.status_code == 200
    assert res.json()["status"] == "Выполнено"


def test_task_update_sets_analytics_fields(client):
    """PATCH выставляет срезы аналитики (приоритет/локация/объект/тема), п. 4.5.1."""
    task_id = _make_task()
    res = client.patch(
        f"/api/tasks/{task_id}",
        json={"priority": "Высокий", "location": "Северный округ", "object": "Школа №1", "theme": "Ремонт"},
    )
    assert res.status_code == 200
    data = res.json()
    assert data["priority"] == "Высокий"
    assert data["location"] == "Северный округ"
    assert data["object"] == "Школа №1"
    assert data["theme"] == "Ремонт"


def test_task_update_rejects_invalid_priority(client):
    """Произвольная строка приоритета отклоняется схемой (контракт API)."""
    task_id = _make_task()
    res = client.patch(f"/api/tasks/{task_id}", json={"priority": "Срочно"})
    assert res.status_code == 422


def test_new_task_defaults_to_normal_priority(client):
    """Задача без явного приоритета получает «Обычный» (дефолт модели)."""
    task_id = _make_task()
    res = client.get(f"/api/tasks/{task_id}")
    assert res.status_code == 200
    assert res.json()["priority"] == "Обычный"


def test_apply_dify_protocol_normalizes_priority():
    """Извлечение протокола мягко приводит приоритет из Dify к TASK_PRIORITIES."""
    from app.api.protocols import _apply_dify_protocol

    with SessionLocal() as db:
        protocol = Protocol(title="P")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        _apply_dify_protocol(
            db,
            protocol,
            {},
            '{"tasks": [{"assignment": "X", "priority": "high", "location": "Центр"},'
            ' {"assignment": "Y", "priority": "чепуха"}]}',
        )
        db.commit()
        tasks = {t.assignment: t for t in db.query(Task).filter(Task.protocol_id == protocol.id)}
        assert tasks["X"].priority == "Высокий"
        assert tasks["X"].location == "Центр"
        # Нераспознанный приоритет -> «Обычный» (мягкая деградация, без ошибки).
        assert tasks["Y"].priority == "Обычный"


def test_apply_dify_protocol_normalizes_status():
    """Статус из Dify приводится к TASK_STATUSES; мусор/пусто -> «Новое»."""
    from app.api.protocols import _apply_dify_protocol

    with SessionLocal() as db:
        protocol = Protocol(title="S")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        _apply_dify_protocol(
            db,
            protocol,
            {},
            '{"tasks": [{"assignment": "A", "status": "Требует проверки"},'
            ' {"assignment": "B", "status": "in progress"},'
            ' {"assignment": "C"}]}',
        )
        db.commit()
        tasks = {t.assignment: t for t in db.query(Task).filter(Task.protocol_id == protocol.id)}
        # Валидный статус workflow сохраняется как есть.
        assert tasks["A"].status == "Требует проверки"
        # Неизвестный/отсутствующий статус -> «Новое» (мягкая деградация).
        assert tasks["B"].status == "Новое"
        assert tasks["C"].status == "Новое"


def test_strip_unresolved_dify_placeholder():
    """Утечка неразрешённого шаблона Dify ({{#context#}}) не показывается пользователю."""
    from app.api.tasks import _strip_unresolved

    assert _strip_unresolved("Контекст отсутствует ({{#context#}})") == ""
    assert _strip_unresolved("Планирование графиков работ") == "Планирование графиков работ"
    assert _strip_unresolved(None) == ""


def test_generate_protocol_dify_error_returns_502_and_no_leftover(client, monkeypatch):
    from app.services import dify_client

    tr = client.post(
        "/api/transcriptions",
        files={"file": ("m.mp3", b"\x00\x01", "audio/mpeg")},
    ).json()

    async def boom(*args, **kwargs):
        return dify_client.DifyResult(answer="", raw={"error": "Dify unavailable"})

    monkeypatch.setattr(dify_client, "run_command", boom)

    before = len(client.get("/api/protocols").json())
    res = client.post("/api/protocols", json={"transcription_id": tr["id"]})
    assert res.status_code == 502
    # Пустой протокол-«пустышка» не должен остаться в БД.
    assert len(client.get("/api/protocols").json()) == before


def test_cannot_demote_or_deactivate_last_admin(client):
    admin = client.post(
        "/api/auth/users",
        json={"username": "boss", "password": "x", "full_name": "Boss", "role": "admin"},
    ).json()

    # Понижение роли последнего администратора запрещено.
    res = client.patch(f"/api/auth/users/{admin['id']}", json={"role": "staff"})
    assert res.status_code == 400

    # Деактивация последнего администратора запрещена.
    res = client.patch(f"/api/auth/users/{admin['id']}", json={"is_active": False})
    assert res.status_code == 400

    # Со вторым администратором понижение первого снова разрешено.
    client.post("/api/auth/users", json={"username": "boss2", "password": "x", "role": "admin"})
    res = client.patch(f"/api/auth/users/{admin['id']}", json={"role": "staff"})
    assert res.status_code == 200
    assert res.json()["role"] == "staff"


@pytest.mark.anyio
async def test_manager_approval_sets_closed_at_in_utc(client):
    """MAX-путь закрытия задачи пишет closed_at в наивном UTC (как и веб-путь)."""
    from app.services import max_handler

    with SessionLocal() as db:
        protocol = Protocol(title="TZ")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        task = Task(protocol_id=protocol.id, assignment="A", completion_text="сделано")
        db.add(task)
        db.commit()
        db.refresh(task)
        task_id = task.id

    with SessionLocal() as db:
        await max_handler._handle_manager_approval(db, {"callback": {}}, f"approve:{task_id}")
        task = db.get(Task, task_id)
        assert task.status == "Выполнено"
        utc_now = datetime.now(timezone.utc).replace(tzinfo=None)
        # UTC, а не Europe/Moscow (иначе расхождение было бы ~3 часа = 10800 с).
        assert abs((utc_now - task.closed_at).total_seconds()) < 120


@pytest.mark.anyio
async def test_reminder_marks_sent_without_touching_notified_at(client, monkeypatch):
    """Напоминание ставит reminder_sent, но не трогает notified_at (см. 1.15)."""
    from app.services import reminders

    with SessionLocal() as db:
        protocol = Protocol(title="R")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        task = Task(protocol_id=protocol.id, assignment="A", deadline="2020-01-01 10:00")
        db.add(task)
        db.commit()
        db.refresh(task)
        task_id = task.id

    async def fake_send(_task):
        return True

    monkeypatch.setattr(reminders, "_send_reminder", fake_send)
    await reminders.scan_once()

    with SessionLocal() as db:
        task = db.get(Task, task_id)
        assert task.reminder_sent is True
        assert task.notified_at is None


def test_send_task_to_max_disabled(client):
    with SessionLocal() as db:
        protocol = Protocol(title="Тест")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        task = Task(protocol_id=protocol.id, assignment="Проверить отправку")
        db.add(task)
        db.commit()
        db.refresh(task)
        task_id = task.id

    res = client.post(f"/api/tasks/{task_id}/send-max")
    assert res.status_code == 400
    assert "MAX не настроен" in res.json()["detail"]


def _make_task(assignment: str = "A", deadline: str = "") -> str:
    with SessionLocal() as db:
        protocol = Protocol(title="P")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        task = Task(protocol_id=protocol.id, assignment=assignment, deadline=deadline)
        db.add(task)
        db.commit()
        db.refresh(task)
        return task.id


def test_patch_deadline_populates_deadline_at(client):
    task_id = _make_task()
    # Полдень -> при переводе МСК->UTC день не меняется, тест детерминирован.
    res = client.patch(f"/api/tasks/{task_id}", json={"deadline": "15.06.2030 12:00"})
    assert res.status_code == 200
    assert res.json()["deadline_at"] is not None
    assert res.json()["deadline_at"].startswith("2030-06-15")


def test_patch_unrecognized_deadline_clears_deadline_at(client):
    task_id = _make_task(deadline="10.10.2030 12:00")
    client.patch(f"/api/tasks/{task_id}", json={"deadline": "10.10.2030 12:00"})
    res = client.patch(f"/api/tasks/{task_id}", json={"deadline": "как получится"})
    assert res.status_code == 200
    assert res.json()["deadline_at"] is None


def test_patch_without_deadline_keeps_deadline_at(client):
    task_id = _make_task()
    client.patch(f"/api/tasks/{task_id}", json={"deadline": "20.07.2030 12:00"})
    # PATCH другого поля не должен затирать вычисленный deadline_at.
    res = client.patch(f"/api/tasks/{task_id}", json={"status": "Требует проверки"})
    assert res.status_code == 200
    assert res.json()["deadline_at"] is not None
    assert res.json()["deadline_at"].startswith("2030-07-20")


def test_qa_session_history_roundtrip(client, monkeypatch):
    from app.services import dify_client

    async def fake_run(*args, **kwargs):
        return dify_client.DifyResult(answer="**Ответ** списком:\n- раз\n- два", raw={})

    monkeypatch.setattr(dify_client, "run_command", fake_run)

    r = client.post(
        "/api/qa",
        json={"question": "Тестовый вопрос", "scope": {"protocol_ids": [], "transcription_ids": []}},
    )
    assert r.status_code == 200
    sid = r.json()["session_id"]

    h = client.get(f"/api/qa/sessions/{sid}")
    assert h.status_code == 200
    data = h.json()
    assert data["session_id"] == sid
    roles = [m["role"] for m in data["messages"]]
    assert roles == ["user", "assistant"]
    assert data["messages"][0]["content"] == "Тестовый вопрос"
    assert "Ответ" in data["messages"][1]["content"]


def test_qa_session_history_404(client):
    assert client.get("/api/qa/sessions/no-such-session").status_code == 404


def test_qa_keeps_context_out_of_routing_query(client, monkeypatch):
    """Маршрутная строка Dify не должна тащить контекст встречи (иначе классификатор
    уводит запрос в ветку извлечения протокола). Контекст идёт через inputs.context."""
    from app.services import dify_client

    captured: dict = {}

    async def fake_run(*args, **kwargs):
        captured.update(kwargs)
        return dify_client.DifyResult(answer="ответ", raw={})

    monkeypatch.setattr(dify_client, "run_command", fake_run)

    # Создаём транскрипт, чтобы в scope был реальный контекст.
    tr = client.post(
        "/api/transcriptions/text",
        json={"title": "Встреча", "text": "[00:00] Иванов: длинный контекст встречи про склад"},
    ).json()
    r = client.post(
        "/api/qa",
        json={"question": "Кто ответственный?", "scope": {"transcription_ids": [tr["id"]], "protocol_ids": []}},
    )
    assert r.status_code == 200
    # В query — только короткий вопрос, без контекста; контекст — в inputs.
    assert captured["query"] == "Кто ответственный?"
    assert "склад" not in captured["query"]
    assert "склад" in captured["inputs"]["context"]


def test_apply_dify_protocol_populates_deadline_at():
    from app.api.protocols import _apply_dify_protocol

    with SessionLocal() as db:
        protocol = Protocol(title="G")
        db.add(protocol)
        db.commit()
        db.refresh(protocol)
        _apply_dify_protocol(
            db,
            protocol,
            {},
            '{"tasks": [{"assignment": "X", "deadline": "05.05.2030 12:00"}]}',
        )
        db.commit()
        task = db.query(Task).filter(Task.protocol_id == protocol.id).one()
        assert task.deadline == "05.05.2030 12:00"
        assert task.deadline_at is not None
        assert task.deadline_at.year == 2030 and task.deadline_at.month == 5


def test_notify_endpoint_delegates_to_send_max(client):
    """Deprecated /api/max/tasks/{id}/notify — тонкая обёртка над /send-max:
    отдаёт ту же 400-ошибку конфигурации MAX (см. 4.5)."""
    task_id = _make_task()
    res = client.post(f"/api/max/tasks/{task_id}/notify")
    assert res.status_code == 400
    assert "MAX не настроен" in res.json()["detail"]


def test_protocol_update(client):
    """PUT /api/protocols/{id}: частичное обновление метаданных/текста протокола.

    Приёмка части 1: обновляет только переданные поля, 404 на несуществующий id,
    а экспорт md отдаёт обновлённый текст (protocol_to_md читает из Protocol)."""
    with SessionLocal() as db:
        p = Protocol(title="Старый заголовок", date="01.01.2030", number="7", body="Старое тело")
        db.add(p)
        db.commit()
        db.refresh(p)
        pid = p.id

    # 404 на несуществующий протокол.
    assert client.put("/api/protocols/no-such-id", json={"title": "X"}).status_code == 404

    # Частичное тело меняет только переданные поля (title/body), остальное — как было.
    res = client.put(f"/api/protocols/{pid}", json={"title": "Новый заголовок", "body": "Новое тело"})
    assert res.status_code == 200
    data = res.json()
    assert data["title"] == "Новый заголовок"
    assert data["body"] == "Новое тело"
    assert data["date"] == "01.01.2030"   # не передавали — не изменилось
    assert data["number"] == "7"

    # Экспорт (md) отдаёт обновлённый текст без дополнительных действий.
    exp = client.post("/api/export", json={"object_type": "protocol", "object_id": pid, "fmt": "md"})
    assert exp.status_code == 200
    assert "Новый заголовок" in exp.text
    assert "Новое тело" in exp.text


def test_protocol_template_upload_detects_placeholders(client):
    """Загрузка .docx: плейсхолдеры распознаются и авто-сопоставляются."""
    lines = [
        "{{ title }} от {{ date }} № {{ number }}",
        "{{ body }}",
        "{% for t in tasks %}{{ t.assignment }}{% endfor %}",
    ]
    up = client.post(
        "/api/protocol-templates",
        files={"file": ("proto.docx", _template_bytes(lines), DOCX_MIME)},
    )
    assert up.status_code == 200
    dto = up.json()
    assert set(dto["detected_placeholders"]) == {"title", "date", "number", "body", "tasks"}
    assert dto["field_mapping"] == {f: f for f in ("title", "date", "number", "body", "tasks")}
    assert dto["is_active"] is True

    # Активный шаблон отдаётся отдельным эндпоинтом.
    active = client.get("/api/protocol-templates/active")
    assert active.status_code == 200
    assert active.json()["id"] == dto["id"]


def test_protocol_template_mapping_update(client):
    """PUT .../mapping меняет маппинг без повторной загрузки файла."""
    up = client.post(
        "/api/protocol-templates",
        files={"file": ("m.docx", _template_bytes(["{{ heading }}"]), DOCX_MIME)},
    )
    assert up.status_code == 200
    tid = up.json()["id"]
    # heading не каноническое имя — авто-маппинг пуст.
    assert up.json()["field_mapping"] == {}

    res = client.put(f"/api/protocol-templates/{tid}/mapping", json={"field_mapping": {"title": "heading"}})
    assert res.status_code == 200
    assert res.json()["field_mapping"] == {"title": "heading"}

    # 404 на несуществующий шаблон.
    assert client.put("/api/protocol-templates/nope/mapping", json={"field_mapping": {}}).status_code == 404


def test_export_protocol_docx_uses_active_template(client):
    """POST /api/export protocol/docx: при активном шаблоне рендерит по нему,
    без активного — прежний generic markdown->docx (обе ветки без 500)."""
    from app.models import ProtocolTemplate

    with SessionLocal() as db:
        p = Protocol(title="Совет директоров", body="Тело протокола")
        db.add(p)
        db.commit()
        db.refresh(p)
        db.add(Task(protocol_id=p.id, assignment="Поручение А", responsible="Петров",
                    deadline="завтра", status="Новое"))
        # Детерминированно снимаем активность со всех шаблонов -> generic-ветка.
        db.query(ProtocolTemplate).update({ProtocolTemplate.is_active: False})
        db.commit()
        pid = p.id

    # Без активного шаблона: работает и НЕ содержит маркер шаблона.
    res = client.post("/api/export", json={"object_type": "protocol", "object_id": pid, "fmt": "docx"})
    assert res.status_code == 200
    assert "ШАБЛОН-МАРКЕР" not in _docx_text_from_bytes(res.content)

    # Загружаем активный шаблон с уникальным маркером.
    lines = ["ШАБЛОН-МАРКЕР {{ title }}", "{% for t in tasks %}{{ t.assignment }}{% endfor %}"]
    up = client.post("/api/protocol-templates", files={"file": ("tpl.docx", _template_bytes(lines), DOCX_MIME)})
    assert up.status_code == 200

    # Теперь docx рендерится по шаблону: маркер + реальные данные протокола.
    res2 = client.post("/api/export", json={"object_type": "protocol", "object_id": pid, "fmt": "docx"})
    assert res2.status_code == 200
    text = _docx_text_from_bytes(res2.content)
    assert "ШАБЛОН-МАРКЕР" in text
    assert "Совет директоров" in text
    assert "Поручение А" in text


def test_startup_cleanup_purges_expired_sessions_and_empty_chats(client):
    """Стартовая чистка (4.8): удаляет истёкшие сессии и старые пустые чаты,
    не трогая живые сессии, старые непустые и свежие пустые чаты."""
    from datetime import timedelta

    from app.models import AuthSession, ChatMessage, ChatSession, User, _now
    from app.services import maintenance

    with SessionLocal() as db:
        user = User(username="cleanup-user", role="admin")
        db.add(user)
        db.commit()
        db.refresh(user)
        db.add_all([
            AuthSession(token="cleanup-expired", user_id=user.id, expires_at=_now() - timedelta(hours=1)),
            AuthSession(token="cleanup-fresh", user_id=user.id, expires_at=_now() + timedelta(hours=1)),
            ChatSession(id="cleanup-old-empty", created_at=_now() - timedelta(days=40)),
            ChatSession(id="cleanup-old-used", created_at=_now() - timedelta(days=40)),
            ChatSession(id="cleanup-recent-empty", created_at=_now()),
        ])
        db.add(ChatMessage(session_id="cleanup-old-used", role="user", content="hi"))
        db.commit()

    with SessionLocal() as db:
        maintenance.run_startup_cleanup(db)

    with SessionLocal() as db:
        assert db.get(AuthSession, "cleanup-expired") is None      # истёкшая — удалена
        assert db.get(AuthSession, "cleanup-fresh") is not None    # живая — сохранена
        assert db.get(ChatSession, "cleanup-old-empty") is None    # старый пустой — удалён
        assert db.get(ChatSession, "cleanup-old-used") is not None  # старый непустой — сохранён
        assert db.get(ChatSession, "cleanup-recent-empty") is not None  # свежий — сохранён
