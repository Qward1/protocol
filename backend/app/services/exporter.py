"""Экспорт объектов (протокол, транскрипт, чат, справка) в разные форматы.

Поддерживаются: docx, pdf, md, txt, json.
Возвращает путь к файлу в storage/exports.
"""

from __future__ import annotations

import json
from pathlib import Path

from app.config import settings
from app.services import storage

SUPPORTED = {"docx", "pdf", "md", "txt", "json"}

# Кандидаты системных TTF с кириллицей (для PDF). Первый существующий — выигрывает.
_FONT_CANDIDATES = [
    "C:/Windows/Fonts/arial.ttf",
    "C:/Windows/Fonts/segoeui.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/Library/Fonts/Arial.ttf",
]
_PDF_FONT_NAME = "DOFont"


# --- сборка markdown-представления для каждого типа объекта ---

def transcription_to_md(t) -> str:
    lines = [f"# Транскрипция: {t.filename}", ""]
    speaker_map = t.speaker_map
    for seg in t.segments:
        ts = _fmt_ts(seg.start)
        name = speaker_map.get(seg.speaker, seg.speaker) if seg.speaker else ""
        speaker = f"**{name}**: " if name else ""
        lines.append(f"`[{ts}]` {speaker}{seg.text}")
    return "\n\n".join(lines)


def protocol_to_md(p) -> str:
    lines = [f"# {p.title or 'Протокол'}", ""]
    if p.date:
        lines.append(f"**Дата:** {p.date}")
    if p.number:
        lines.append(f"**Номер:** {p.number}")
    lines.append("")
    if p.body:
        lines.append(p.body)
        lines.append("")
    if p.tasks:
        lines.append("## Поручения")
        lines.append("")
        lines.append("| Поручение | Ответственный | Срок | Статус |")
        lines.append("|---|---|---|---|")
        for task in p.tasks:
            lines.append(
                f"| {task.assignment} | {task.responsible} | {task.deadline} | {task.status} |"
            )
    return "\n".join(lines)


def chat_to_md(session) -> str:
    lines = [f"# {session.title}", ""]
    for msg in session.messages:
        who = "Вопрос" if msg.role == "user" else "Ответ"
        lines.append(f"**{who}:** {msg.content}")
        citations = json.loads(msg.citations_json or "[]")
        for cit in citations:
            lines.append(f"> Источник: {cit.get('title','')} — {cit.get('fragment','')}")
        lines.append("")
    return "\n".join(lines)


def justification_to_md(j, task) -> str:
    return "\n".join(
        [
            "# Справка-обоснование назначения поручения",
            "",
            f"**Поручение:** {task.assignment}",
            f"**Ответственный:** {task.responsible}",
            "",
            "## На основании фрагмента записи",
            j.fragment or task.source_fragment,
            "",
            "## На основании должностной обязанности",
            j.duty,
            "",
            "## Обоснование",
            j.text,
        ]
    )


def _fmt_ts(seconds: float) -> str:
    s = int(seconds)
    return f"{s // 60:02d}:{s % 60:02d}"


# --- рендереры форматов ---

def _write_txt(md: str, out: Path) -> Path:
    out.write_text(md, encoding="utf-8")
    return out


def _write_md(md: str, out: Path) -> Path:
    out.write_text(md, encoding="utf-8")
    return out


def _write_json(obj: dict, out: Path) -> Path:
    out.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def _parse_md_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _is_md_separator(line: str) -> bool:
    cells = _parse_md_row(line)
    return bool(cells) and all(set(c) <= {"-", ":", " "} and "-" in c for c in cells)


def _write_docx(md: str, out: Path) -> Path:
    from docx import Document  # python-docx

    doc = Document()
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        # Блок markdown-таблицы -> настоящая таблица python-docx.
        if line.lstrip().startswith("|") and i + 1 < len(lines) and _is_md_separator(lines[i + 1]):
            header = _parse_md_row(line)
            rows: list[list[str]] = []
            i += 2  # пропускаем заголовок и разделитель
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(_parse_md_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, text in enumerate(header):
                table.rows[0].cells[j].text = text
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    cells[j].text = row[j] if j < len(row) else ""
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            doc.add_paragraph(line.replace("**", ""))
        else:
            doc.add_paragraph("")
        i += 1
    doc.save(str(out))
    return out


def _register_pdf_font() -> str:
    """Зарегистрировать TTF с кириллицей. Возвращает имя шрифта (или Helvetica)."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if _PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return _PDF_FONT_NAME
    candidates = [settings.export.pdf_font_path, *_FONT_CANDIDATES]
    for path in candidates:
        if path and Path(path).is_file():
            try:
                pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, path))
                return _PDF_FONT_NAME
            except Exception:
                continue
    return "Helvetica"  # фоллбэк (кириллица может не отобразиться)


def _write_pdf(md: str, out: Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    font = _register_pdf_font()
    doc = SimpleDocTemplate(str(out), pagesize=A4)
    styles = getSampleStyleSheet()
    # Применяем кириллический шрифт ко всем используемым стилям.
    for name in ("BodyText", "Heading1", "Heading2"):
        styles[name].fontName = font

    story = []
    for line in md.splitlines():
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
        style = styles["Heading1"] if line.startswith("# ") else (
            styles["Heading2"] if line.startswith("## ") else styles["BodyText"]
        )
        text = escape(line.lstrip("# ").replace("**", ""))
        story.append(Paragraph(text, style))
    doc.build(story)
    return out


def render(md: str, raw_obj: dict, fmt: str, name: str) -> Path:
    """Сохранить объект в формате fmt. raw_obj используется для json."""
    if fmt not in SUPPORTED:
        raise ValueError(f"Unsupported format: {fmt}")
    out = storage.exports_dir() / f"{storage.safe_name(name)}.{fmt}"
    if fmt == "txt":
        return _write_txt(md, out)
    if fmt == "md":
        return _write_md(md, out)
    if fmt == "json":
        return _write_json(raw_obj, out)
    if fmt == "docx":
        return _write_docx(md, out)
    if fmt == "pdf":
        return _write_pdf(md, out)
    raise ValueError(fmt)
